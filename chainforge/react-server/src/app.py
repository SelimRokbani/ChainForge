import os
from flask import Flask, request, jsonify
from flask_cors import CORS
import nltk
import spacy
import io
import fitz
from docx import Document
import cohere
import gensim
from gensim.utils import simple_preprocess
from gensim.corpora import Dictionary
from sklearn.cluster import KMeans
import tiktoken
from langchain.text_splitter import RecursiveCharacterTextSplitter
from bertopic import BERTopic
from dotenv import load_dotenv
import logging
import math
import tempfile
import openai

# For Retrieval Methods
from sklearn.feature_extraction.text import TfidfVectorizer
from rank_bm25 import BM25Okapi
from whoosh import index
from whoosh.fields import Schema, TEXT, ID
from whoosh.qparser import QueryParser

# Add these imports at the top
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
import faiss
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_community.embeddings import HuggingFaceEmbeddings

load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

nltk.download("punkt")
try:
    nlp = spacy.load("en_core_web_sm")
    logger.info("Loaded spaCy model 'en_core_web_sm'.")
except OSError:
    from spacy.cli import download
    download("en_core_web_sm")
    nlp = spacy.load("en_core_web_sm")
    logger.info("Downloaded and loaded spaCy model 'en_core_web_sm'.")

COHERE_API_KEY = os.getenv("COHERE_API_KEY", "")
if COHERE_API_KEY:
    co = cohere.Client(COHERE_API_KEY)
    logger.info("Initialized Cohere client.")
else:
    co = None
    logger.warning("Cohere API key not found. Cohere-based chunking methods unavailable.")

openai.api_key = os.getenv("OPENAI_API_KEY", "")
if not openai.api_key:
    logger.warning("OpenAI API key not found. OpenAI-based embeddings will fail unless you set OPENAI_API_KEY.")

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

# -------------------- File Extraction Helpers --------------------
def extract_text_from_pdf(file_bytes):
    try:
        text = ""
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for page in doc:
                text += page.get_text()
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise

def extract_text_from_docx(file_bytes):
    try:
        file_stream = io.BytesIO(file_bytes)
        doc = Document(file_stream)
        full_text = [para.text for para in doc.paragraphs]
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise

def extract_text_from_txt(file_bytes):
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {e}")
        raise

# -------------------- Chunking Methods --------------------
from transformers import AutoTokenizer
huggingface_tokenizer = AutoTokenizer.from_pretrained("bert-base-uncased")

def overlapping_langchain_textsplitter(text, chunk_size=200, chunk_overlap=50, keep_separator=True):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size, chunk_overlap=chunk_overlap, keep_separator=keep_separator
    )
    chunks = splitter.split_text(text)
    return chunks if chunks else [text]

def overlapping_openai_tiktoken(text, chunk_size=200, chunk_overlap=50):
    import tiktoken
    enc = tiktoken.encoding_for_model("gpt-3.5-turbo")
    tokens = enc.encode(text)
    result = []
    start = 0
    while start < len(tokens):
        end = start + chunk_size
        chunk = enc.decode(tokens[start:end])
        result.append(chunk)
        start = end - chunk_overlap
        if start < 0:
            start = 0
        if end >= len(tokens):
            break
    return result if result else [text]

def overlapping_huggingface_tokenizers(text, chunk_size=200, chunk_overlap=50):
    tokens = huggingface_tokenizer.encode(text)
    result = []
    start = 0
    while start < len(tokens):
        end = start + chunk_size
        chunk = huggingface_tokenizer.decode(tokens[start:end], skip_special_tokens=True)
        result.append(chunk)
        start = end - chunk_overlap
        if start < 0:
            start = 0
        if end >= len(tokens):
            break
    return result if result else [text]

def syntax_spacy(text):
    doc = nlp.pipe([text])
    sents = []
    for d in doc:
        for s in d.sents:
            sent_text = s.text.strip()
            if sent_text:
                sents.append(sent_text)
    return sents if sents else [text]

def syntax_texttiling(text):
    from nltk.tokenize import TextTilingTokenizer
    tt = TextTilingTokenizer()
    chunks = tt.tokenize(text)
    return chunks if chunks else [text]

def hybrid_texttiling_spacy(text, chunk_size=300, chunk_overlap=100):
    from nltk.tokenize import TextTilingTokenizer
    tt = TextTilingTokenizer()
    sem_chunks = tt.tokenize(text) or [text]
    final_chunks = []
    for chunk in sem_chunks:
        doc = nlp(chunk)
        sentences = [s.text.strip() for s in doc.sents if s.text.strip()]
        buf = []
        current_size = 0
        for sentence in sentences:
            if current_size + len(sentence) > chunk_size:
                final_chunks.append(" ".join(buf))
                buf = buf[-chunk_overlap:] + [sentence]
                current_size = sum(len(x) for x in buf)
            else:
                buf.append(sentence)
                current_size += len(sentence)
        if buf:
            final_chunks.append(" ".join(buf))
    return final_chunks if final_chunks else [text]

def semantic_bertopic(text, n_topics=2, min_topic_size=1):
    topic_model = BERTopic(n_topics=n_topics, min_topic_size=min_topic_size, random_state=42)
    sentences = nltk.sent_tokenize(text)
    if not sentences:
        return [text]
    topics = topic_model.fit_transform(sentences)
    clusters = {}
    for t, s in zip(topics, sentences):
        clusters.setdefault(t, []).append(s)
    return [" ".join(v) for v in clusters.values()] or [text]

def hybrid_bertopic_spacy(text, min_topic_size=2):
    sentences = [s.text.strip() for s in nlp(text).sents if s.text.strip()]
    if not sentences:
        return [text]
    topic_model = BERTopic(min_topic_size=min_topic_size, random_state=42)
    topics = topic_model.fit_transform(sentences)
    clusters = {}
    for t, s in zip(topics, sentences):
        clusters.setdefault(t, []).append(s)
    return [" ".join(v) for v in clusters.values()] or [text]

def hybrid_recursive_gensim(text, chunk_size=300, chunk_overlap=100):
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    sem_chunks = splitter.split_text(text) or [text]
    final_chunks = []
    for sc in sem_chunks:
        sents = nltk.sent_tokenize(sc)
        processed = [simple_preprocess(s) for s in sents]
        dictionary = Dictionary(processed)
        if len(dictionary) == 0:
            continue
        corpus = [dictionary.doc2bow(d) for d in processed]
        from gensim.models import LdaModel
        lda = LdaModel(corpus, id2word=dictionary, num_topics=2, passes=10, random_state=42)
        assignments = []
        for bow in corpus:
            topic_probs = lda.get_document_topics(bow)
            if topic_probs:
                dominant_topic = max(topic_probs, key=lambda x: x[1])[0]
            else:
                dominant_topic = 0
            assignments.append(dominant_topic)
        clusters = {}
        for lbl, st in zip(assignments, sents):
            clusters.setdefault(lbl, []).append(st)
        final_chunks.extend([" ".join(val) for val in clusters.values()])
    return final_chunks if final_chunks else [text]

def hybrid_recursive_cohere(text, chunk_size=300, chunk_overlap=100):
    if not co:
        return [text]
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    sem_chunks = splitter.split_text(text) or [text]
    final_chunks = []
    for sc in sem_chunks:
        sents = nltk.sent_tokenize(sc)
        if not sents:
            continue
        embeddings = co.embed(texts=sents).embeddings
        kmeans = KMeans(n_clusters=2, random_state=42)
        labels = kmeans.fit_predict(embeddings)
        clusters = {}
        for lbl, st in zip(labels, sents):
            clusters.setdefault(lbl, []).append(st)
        final_chunks.extend([" ".join(val) for val in clusters.values()])
    return final_chunks if final_chunks else [text]

def hybrid_recursive_bertopic(text, chunk_size=300, chunk_overlap=300, n_topics=2, min_topic_size=1):
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    sem_chunks = splitter.split_text(text) or [text]
    final_chunks = []
    for sc in sem_chunks:
        sents = nltk.sent_tokenize(sc)
        if not sents:
            continue
        topic_model = BERTopic(n_topics=n_topics, min_topic_size=min_topic_size, random_state=42)
        topics, probs = topic_model.fit_transform(sents)
        clusters = {}
        for t, s0 in zip(topics, sents):
            clusters.setdefault(t, []).append(s0)
        final_chunks.extend([" ".join(val) for val in clusters.values()])
    return final_chunks if final_chunks else [text]

CHUNKING_METHODS = {
    ("Overlapping Chunking", "LangChain's TextSplitter"): overlapping_langchain_textsplitter,
    ("Overlapping Chunking", "OpenAI tiktoken"): overlapping_openai_tiktoken,
    ("Overlapping Chunking", "HuggingFace Tokenizers"): overlapping_huggingface_tokenizers,
    ("Syntax-Based Chunking", "spaCy Sentence Splitter"): syntax_spacy,
    ("Syntax-Based Chunking", "TextTilingTokenizer"): syntax_texttiling,
    ("Hybrid Chunking", "TextTiling + spaCy"): hybrid_texttiling_spacy,
    ("Hybrid Chunking", "BERTopic + spaCy"): hybrid_bertopic_spacy,
    ("Hybrid Chunking", "Recursive TextSplitter + Gensim"): hybrid_recursive_gensim,
    ("Hybrid Chunking", "Recursive TextSplitter + Cohere"): hybrid_recursive_cohere,
    ("Hybrid Chunking", "Recursive TextSplitter + BERTopic"): hybrid_recursive_bertopic,
}

def chunk_text(method_name, library, text, **kwargs):
    func = CHUNKING_METHODS.get((method_name, library))
    if not func:
        logger.warning(f"No chunking function for method='{method_name}', library='{library}'. Returning full text.")
        return [text]
    try:
        chunks = func(text, **kwargs)
        return chunks
    except TypeError:
        # Fallback if the function doesn't accept all kwargs
        return func(text)

# -------------------- Vector & Similarity Helpers --------------------
def cosine_similarity(a, b):
    dot = sum(x * y for x, y in zip(a, b))
    normA = math.sqrt(sum(x * x for x in a))
    normB = math.sqrt(sum(y * y for y in b))
    return dot / (normA * normB + 1e-10)    

# -------------- Whoosh (Optional) --------------
whoosh_index_dir = tempfile.mkdtemp()
schema = Schema(id=ID(stored=True, unique=True), content=TEXT(stored=True))
if not os.path.exists(whoosh_index_dir):
    os.mkdir(whoosh_index_dir)
ix = index.create_in(whoosh_index_dir, schema)

# -------------- Flask Routes --------------

@app.route("/upload", methods=["POST"])
def upload_files():
    """Extract text from the uploaded file based on its extension."""
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file uploaded"}), 400

    filename = file.filename.lower()
    ext = filename.rsplit(".", 1)[-1]
    file_bytes = file.read()
    if ext == "pdf":
        text = extract_text_from_pdf(file_bytes)
    elif ext == "txt":
        text = extract_text_from_txt(file_bytes)
    elif ext == "docx":
        text = extract_text_from_docx(file_bytes)
    else:
        return jsonify({"error": f"Unsupported file type: {ext}"}), 400

    return jsonify({"text": text}), 200

@app.route("/process", methods=["POST"])
def process_file():
    """
    Extract and chunk text using specified chunking method (methodName + library).
    Either a file or raw text can be provided.
    """
    file = request.files.get("file")
    method_name = request.form.get("methodName")
    library = request.form.get("library")
    text_input = request.form.get("text", None)

    settings = {}
    for k in request.form:
        if k not in ["methodName", "library", "text", "file"]:
            val = request.form.get(k)
            if k in ["chunk_size", "chunk_overlap", "n_topics", "min_topic_size"]:
                try:
                    val = int(val)
                except ValueError:
                    val = None
            elif k in ["keep_separator"]:
                val = (val.lower() == "true")
            settings[k] = val

    if not method_name or not library:
        return jsonify({"error": "Missing methodName or library"}), 400

    # Extract text
    if file:
        filename = file.filename.lower()
        ext = filename.rsplit(".", 1)[-1]
        file_bytes = file.read()
        if ext == "pdf":
            text = extract_text_from_pdf(file_bytes)
        elif ext == "txt":
            text = extract_text_from_txt(file_bytes)
        elif ext == "docx":
            text = extract_text_from_docx(file_bytes)
        else:
            return jsonify({"error": f"Unsupported file type: {ext}"}), 400
    elif text_input:
        text = text_input
    else:
        return jsonify({"error": "No file or text provided"}), 400

    # Chunk
    chunks = chunk_text(method_name, library, text, **{k: v for k, v in settings.items() if v is not None})
    return jsonify({
        "library": library,
        "settings": settings,
        "chunks": chunks
    }), 200

class EmbeddingModelRegistry:
    _models = {}
    
    @classmethod
    def register(cls, model_name):
        def decorator(embedding_func):
            cls._models[model_name] = embedding_func
            return embedding_func
        return decorator
        
    @classmethod
    def get_embedder(cls, model_name):
        return cls._models.get(model_name)
    
    @classmethod
    def list_models(cls):
        return list(cls._models.keys())

@EmbeddingModelRegistry.register("huggingface")
def huggingface_embedder(texts, model_name="sentence-transformers/all-mpnet-base-v2"):
    """
    Generate embeddings using HuggingFace Transformers.
    
    Args:
        texts: List of text strings to embed
        model_name: HuggingFace model name/path (default: sentence-transformers/all-mpnet-base-v2)
        
    Returns:
        List of embeddings for each text
    """
    try:
        from transformers import AutoTokenizer, AutoModel
        import torch
        
        logger.info(f"Using HuggingFace model: {model_name} for {len(texts)} texts")
        
        # Load model and tokenizer
        tokenizer = AutoTokenizer.from_pretrained(model_name)
        model = AutoModel.from_pretrained(model_name)
        
        embeddings = []
        batch_size = 32
        for i in range(0, len(texts), batch_size):
            batch_texts = texts[i:i+batch_size]
            batch_embeddings = []
            
            for t in batch_texts:
                inputs = tokenizer(t, return_tensors="pt", truncation=True, padding=True, 
                                  max_length=512)  # Add max_length for safety
                with torch.no_grad():
                    outputs = model(**inputs)
                # Use mean pooling by default
                emb = outputs.last_hidden_state.mean(dim=1).squeeze().tolist()
                batch_embeddings.append(emb)
                
            embeddings.extend(batch_embeddings)
            
        return embeddings
    except Exception as e:
        logger.error(f"HuggingFace embedder failed: {str(e)}")
        raise ValueError(f"Failed to generate HuggingFace embeddings: {str(e)}")

@EmbeddingModelRegistry.register("OpenAI Embeddings")
def openai_embedder(texts, model_name="text-embedding-ada-002"):
    """
    Generate embeddings using OpenAI Embeddings.
    
    Args:
        texts: List of text strings to embed
        model_name: OpenAI embedding model to use (default: text-embedding-ada-002)
        
    Returns:
        List of embeddings for each text
    """
    try:
        import openai
        logger.info(f"Using OpenAI model: {model_name} for {len(texts)} texts")
        
        embeddings = []
        # Process in batches of 16 to stay within rate limits
        batch_size = 16
        for i in range(0, len(texts), batch_size):
            batch_texts = texts[i:i+batch_size]
            batch_embeddings = []
            
            for t in batch_texts:
                resp = openai.Embedding.create(input=t, model=model_name)
                emb = resp["data"][0]["embedding"]
                batch_embeddings.append(emb)
                
            embeddings.extend(batch_embeddings)
            
        return embeddings
    except Exception as e:
        logger.error(f"OpenAI embedder failed: {str(e)}")
        raise ValueError(f"Failed to generate OpenAI embeddings: {str(e)}")

@EmbeddingModelRegistry.register("Cohere Embeddings")
def cohere_embedder(texts, model_name="embed-english-v2.0"):
    """
    Generate embeddings using Cohere Embeddings.
    
    Args:
        texts: List of text strings to embed
        model_name: Cohere embedding model to use (default: embed-english-v2.0)
        
    Returns:
        List of embeddings for each text
    """
    try:
        import cohere
        logger.info(f"Using Cohere model: {model_name} for {len(texts)} texts")
        
        # Get API key from environment or settings
        api_key = os.environ.get("COHERE_API_KEY")
        if not api_key:
            from flask import current_app
            api_key = current_app.config.get("COHERE_API_KEY")
            
        if not api_key:
            raise ValueError("Cohere API key not found in environment or app config")
            
        co = cohere.Client(api_key)
        
        batch_size = 32  
        embeddings = []
        
        for i in range(0, len(texts), batch_size):
            batch_texts = texts[i:i+batch_size]
            response = co.embed(texts=batch_texts, model=model_name)
            embeddings.extend(response.embeddings)
            
        return embeddings
    except Exception as e:
        logger.error(f"Cohere embedder failed: {str(e)}")
        raise ValueError(f"Failed to generate Cohere embeddings: {str(e)}")

@EmbeddingModelRegistry.register("Sentence Transformers")
def sentence_transformers_embedder(texts, model_name="all-MiniLM-L6-v2"):
    """
    Generate embeddings using Sentence Transformers.
    
    Args:
        texts: List of text strings to embed
        model_name: Sentence Transformers model name (default: all-MiniLM-L6-v2)
        
    Returns:
        List of embeddings for each text
    """
    try:
        from sentence_transformers import SentenceTransformer
        logger.info(f"Using SentenceTransformer model: {model_name} for {len(texts)} texts")
        
        model = SentenceTransformer(model_name)
        
        # Process in reasonable batch sizes
        batch_size = 32
        embeddings = []
        
        for i in range(0, len(texts), batch_size):
            batch_texts = texts[i:i+batch_size]
            embeddings = model.encode(batch_texts).tolist()
            embeddings.extend(embeddings)
            
        return embeddings
    except Exception as e:
        logger.error(f"SentenceTransformer embedder failed: {str(e)}")
        raise ValueError(f"Failed to generate SentenceTransformer embeddings: {str(e)}")

# Define a registry for retrieval methods
class RetrievalMethodRegistry:
    _methods = {}
    
    @classmethod
    def register(cls, method_name):
        def decorator(handler_func):
            cls._methods[method_name] = handler_func
            return handler_func
        return decorator
        
    @classmethod
    def get_handler(cls, method_name):
        return cls._methods.get(method_name)

@RetrievalMethodRegistry.register("bm25")
def handle_bm25(chunk_objs, queries, settings):
    top_k = settings.get("top_k", 5)
    k1 = settings.get("bm25_k1", 1.5)
    b = settings.get("bm25_b", 0.75)
    
    # Extract text from chunk objects
    chunk_texts = [chunk.get("text", "") for chunk in chunk_objs]
    
    # Preprocess corpus once
    tokenized_corpus = [simple_preprocess(doc) for doc in chunk_texts]
    bm25 = BM25Okapi(tokenized_corpus, k1=k1, b=b)
    
    results = {}
    for query in queries:
        tokenized_query = simple_preprocess(query)
        raw_scores = bm25.get_scores(tokenized_query)
        
        # Normalize scores
        max_score = max(raw_scores) if raw_scores.any() and max(raw_scores) > 0 else 1
        normalized_scores = [score / max_score for score in raw_scores]
        
        # Build result objects with all the necessary metadata
        retrieved = []
        scored_chunks = sorted(zip(chunk_objs, normalized_scores), key=lambda x: x[1], reverse=True)
        
        for chunk, similarity in scored_chunks[:top_k]:
            retrieved.append({
                "text": chunk.get("text", ""),
                "similarity": float(similarity),
                "docTitle": chunk.get("docTitle", ""),
                "chunkId": chunk.get("chunkId", ""),
            })
        
        results[query] = retrieved
    
    return results

@RetrievalMethodRegistry.register("tfidf")
def handle_tfidf(chunk_objs, queries, settings):
    top_k = settings.get("top_k", 5)
    max_features = settings.get("max_features", 500)
    
    # Extract text from chunk objects
    chunk_texts = [chunk.get("text", "") for chunk in chunk_objs]
    
    # Create and fit vectorizer once for all queries
    vectorizer = TfidfVectorizer(stop_words="english", max_features=max_features)
    tfidf_matrix = vectorizer.fit_transform(chunk_texts)
    
    results = {}
    for query in queries:
        try:
            query_vec = vectorizer.transform([query])
            # Calculate cosine similarities
            sims = (tfidf_matrix * query_vec.T).toarray().flatten()

            # Normalize scores (0 to 1 range relative to max score)
            max_sim = sims.max() if sims.size > 0 and sims.max() > 0 else 1.0
            normalized_sims = sims / max_sim

            # Build result objects
            retrieved = []
            # Get indices of top_k scores using numpy.argsort for efficiency
            ranked_idx = np.argsort(normalized_sims)[::-1][:top_k] # High to low

            for i in ranked_idx:
                # Ensure we don't index out of bounds if corpus was smaller than top_k
                if i < len(chunk_objs):
                    chunk = chunk_objs[i]
                    retrieved.append({
                        "text": chunk.get("text", ""),
                        "similarity": float(normalized_sims[i]), # Ensure standard float
                        "docTitle": chunk.get("docTitle", ""),
                        "chunkId": chunk.get("chunkId", ""),
                    })

            results[query] = retrieved
        except Exception as e:
             logger.error(f"Error processing query '{query}' with TF-IDF: {e}")
             results[query] = [] # Return empty list for this specific query on error

    return results

@RetrievalMethodRegistry.register("boolean")
def handle_boolean(chunk_objs, queries, settings):
    top_k = settings.get("top_k", 5)
    required_match_count = settings.get("required_match_count", 1)
    
    # Extract text from chunk objects
    chunk_texts = [chunk.get("text", "") for chunk in chunk_objs]
    
    results = {}
    for query in queries:
        q_tokens = set(simple_preprocess(query))
        if len(q_tokens) < required_match_count:
            # Not enough tokens in query to match the required count
            results[query] = []
            continue
            
        scored = []
        for i, c in enumerate(chunk_texts):
            c_tokens = set(simple_preprocess(c))
            matches = len(q_tokens.intersection(c_tokens))
            if matches >= required_match_count:
                score = matches / (len(c_tokens) + 1e-9)  # Normalize by document length
                scored.append((i, score))
                
        # Sort by score
        scored.sort(key=lambda x: x[1], reverse=True)
        
        # Normalize scores
        retrieved = []
        if scored:
            max_score = scored[0][1]
            for i, score in scored[:top_k]:
                chunk = chunk_objs[i]
                normalized_score = score / max_score if max_score > 0 else 0
                retrieved.append({
                    "text": chunk.get("text", ""),
                    "similarity": float(normalized_score),
                    "docTitle": chunk.get("docTitle", ""),
                    "chunkId": chunk.get("chunkId", ""),
                })
        
        results[query] = retrieved
    
    return results

@RetrievalMethodRegistry.register("overlap")
def handle_keyword_overlap(chunk_objs, queries, settings):
    top_k = settings.get("top_k", 5)
    
    # Extract text from chunk objects
    chunk_texts = [chunk.get("text", "") for chunk in chunk_objs]
    
    results = {}
    for query in queries:
        q_tokens = set(simple_preprocess(query))
        scored = []
        
        for i, c in enumerate(chunk_texts):
            c_tokens = set(simple_preprocess(c))
            overlap = len(q_tokens.intersection(c_tokens))
            scored.append((i, overlap))
            
        # Sort by overlap count
        scored.sort(key=lambda x: x[1], reverse=True)
        
        # Normalize scores
        retrieved = []
        if scored and scored[0][1] > 0:  # Ensure max score > 0
            max_score = scored[0][1]
            for i, score in scored[:top_k]:
                chunk = chunk_objs[i]
                normalized_score = score / max_score
                retrieved.append({
                    "text": chunk.get("text", ""),
                    "similarity": float(normalized_score),
                    "docTitle": chunk.get("docTitle", ""),
                    "chunkId": chunk.get("chunkId", ""),
                })
        else:
            # No overlaps found
            retrieved = []
        
        results[query] = retrieved
    
    return results

import numpy as np
import heapq
from sklearn.metrics.pairwise import cosine_similarity as sklearn_cosine
from sklearn.cluster import KMeans
import math


# Helper functions for similarity calculations
def cosine_similarity(vec1, vec2):
    """Compute cosine similarity between two vectors"""
    dot_product = sum(a * b for a, b in zip(vec1, vec2))
    norm_a = math.sqrt(sum(a * a for a in vec1))
    norm_b = math.sqrt(sum(b * b for b in vec2))
    return dot_product / (norm_a * norm_b) if norm_a * norm_b > 0 else 0

def manhattan_distance(vec1, vec2):
    """Compute Manhattan distance between two vectors"""
    return sum(abs(a - b) for a, b in zip(vec1, vec2))

def euclidean_distance(vec1, vec2):
    """Compute Euclidean distance between two vectors"""
    return math.sqrt(sum((a - b) ** 2 for a, b in zip(vec1, vec2)))

@RetrievalMethodRegistry.register("cosine")
def handle_cosine_similarity(chunks, chunk_embeddings, queries, query_embeddings, settings):
    """
    Retrieve chunks using cosine similarity between embeddings.
    
    This implementation uses a min-heap to keep only the top-k results in memory.
    """
    top_k = settings.get("top_k", 5)
    results = {}
    
    for (query, query_emb) in zip(queries, query_embeddings):
        # Use a min heap to keep track of top k results
        min_heap = []
        
        # Calculate similarities and maintain heap of size top_k
        for i, (chunk, chunk_emb) in enumerate(zip(chunks, chunk_embeddings)):
            sim = cosine_similarity(chunk_emb, query_emb)
            
            # If heap is not full, add the item
            if len(min_heap) < top_k:
                heapq.heappush(min_heap, (sim, i))
            # If similarity is higher than the smallest in heap, replace it
            elif sim > min_heap[0][0]:
                heapq.heappushpop(min_heap, (sim, i))
        
        # Convert heap to sorted results (highest similarity first)
        retrieved = []
        for sim, i in sorted(min_heap, reverse=True):
            chunk = chunks[i]
            retrieved.append({
                "text": chunk.get("text", ""),
                "similarity": float(sim),
                "docTitle": chunk.get("docTitle", ""),
                "chunkId": chunk.get("chunkId", ""),
            })
        
        results[query] = retrieved
    
    return results

@RetrievalMethodRegistry.register("manhattan")
def handle_manhattan(chunk_objs, chunk_embeddings, queries, query_embeddings, settings):
    """
    Retrieve chunks using Manhattan distance between embeddings.
    """
    top_k = settings.get("top_k", 5)
    results = {}
    
    for query, query_emb in zip(queries, query_embeddings):
        # Use a min heap to keep track of top k results
        min_heap = []
        
        # Calculate similarities and maintain heap of size top_k
        for i, (chunk, chunk_emb) in enumerate(zip(chunk_objs, chunk_embeddings)):
            # Lower Manhattan distance = higher similarity
            distance = manhattan_distance(chunk_emb, query_emb)
            sim = 1.0 / (1.0 + distance)  # Transform to similarity score
            
            if len(min_heap) < top_k:
                heapq.heappush(min_heap, (sim, i))
            elif sim > min_heap[0][0]:
                heapq.heappushpop(min_heap, (sim, i))
        
        # Convert heap to sorted results
        retrieved = []
        for sim, i in sorted(min_heap, reverse=True):
            chunk = chunk_objs[i]
            retrieved.append({
                "text": chunk.get("text", ""),
                "similarity": float(sim),
                "docTitle": chunk.get("docTitle", ""),
                "chunkId": chunk.get("chunkId", ""),
            })
        
        results[query] = retrieved
    
    return results

@RetrievalMethodRegistry.register("euclidean")
def handle_euclidean(chunk_objs, chunk_embeddings, queries, query_embeddings, settings):
    """
    Retrieve chunks using Euclidean distance between embeddings.
    """
    top_k = settings.get("top_k", 5)
    results = {}
    
    for query, query_emb in zip(queries, query_embeddings):
        min_heap = []
        
        for i, (chunk, chunk_emb) in enumerate(zip(chunk_objs, chunk_embeddings)):
            distance = euclidean_distance(chunk_emb, query_emb)
            sim = 1.0 / (1.0 + distance)  # Transform to similarity score
            
            if len(min_heap) < top_k:
                heapq.heappush(min_heap, (sim, i))
            elif sim > min_heap[0][0]:
                heapq.heappushpop(min_heap, (sim, i))
        
        # Convert heap to sorted results
        retrieved = []
        for sim, i in sorted(min_heap, reverse=True):
            chunk = chunk_objs[i]
            retrieved.append({
                "text": chunk.get("text", ""),
                "similarity": float(sim),
                "docTitle": chunk.get("docTitle", ""),
                "chunkId": chunk.get("chunkId", ""),
            })
        
        results[query] = retrieved
    
    return results

@RetrievalMethodRegistry.register("clustered")
def handle_clustered(chunk_objs, chunk_embeddings, queries, query_embeddings, settings):
    """
    Retrieve chunks using a combination of query similarity and cluster similarity.
    """
    top_k = settings.get("top_k", 5)
    n_clusters = settings.get("n_clusters", 3)
    query_coeff = settings.get("query_coeff", 0.6)
    center_coeff = settings.get("center_coeff", 0.4)
    results = {}
    
    # Convert embeddings to numpy array for clustering
    X = np.array(chunk_embeddings)
    
    # Only perform clustering if we have enough samples
    if len(X) >= 2:
        n_clusters = min(n_clusters, len(X))
        kmeans = KMeans(n_clusters=n_clusters, random_state=42)
        labels = kmeans.fit_predict(X)
        cluster_centers = kmeans.cluster_centers_
        
        for query, query_emb in zip(queries, query_embeddings):
            min_heap = []
            query_emb_np = np.array(query_emb).reshape(1, -1)
            
            for i, (chunk, chunk_emb) in enumerate(zip(chunk_objs, chunk_embeddings)):
                # Calculate similarity to query
                chunk_emb_np = np.array(chunk_emb).reshape(1, -1)
                query_sim = float(sklearn_cosine(chunk_emb_np, query_emb_np)[0][0])
                
                # Calculate similarity to cluster center
                center_sim = float(sklearn_cosine(
                    chunk_emb_np, 
                    cluster_centers[labels[i]].reshape(1, -1)
                )[0][0])
                
                # Combined similarity score (weighted)
                combined_sim = query_coeff * query_sim + center_coeff * center_sim
                
                if len(min_heap) < top_k:
                    heapq.heappush(min_heap, (combined_sim, i))
                elif combined_sim > min_heap[0][0]:
                    heapq.heappushpop(min_heap, (combined_sim, i))
            
            # Convert heap to sorted results
            retrieved = []
            for sim, i in sorted(min_heap, reverse=True):
                chunk = chunk_objs[i]
                retrieved.append({
                    "text": chunk.get("text", ""),
                    "similarity": float(sim),
                    "docTitle": chunk.get("docTitle", ""),
                    "chunkId": chunk.get("chunkId", ""),
                })
            
            results[query] = retrieved
    return results

import faiss
import numpy as np
import os
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings # Import base class
from typing import List # For type hinting
# Assuming RetrievalMethodRegistry is defined elsewhere
# Assuming logger is configured elsewhere, but we won't use it here

# --- Define DummyEmbeddings Class ---
class DummyEmbeddings(Embeddings):
    """
    A dummy embedding class implementing the LangChain Embeddings interface.
    Used when pre-computed embeddings are provided to FAISS wrappers.
    Returns zero vectors of the specified dimension.
    """
    def __init__(self, dimension: int):
        self.dimension = dimension
        if not isinstance(dimension, int) or dimension <= 0:
             raise ValueError(f"Dimension must be a positive integer, got {dimension}")

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """Return zero vectors for a list of documents."""
        return [[0.0] * self.dimension for _ in texts]

    def embed_query(self, text: str) -> List[float]:
        """Return a single zero vector for a query."""
        return [0.0] * self.dimension

@RetrievalMethodRegistry.register("faiss")
def handle_faiss(chunk_objs, chunk_embeddings, queries, query_embeddings, settings):
    """
    Retrieve chunks using FAISS with explicit metric control (L2/IP) for MVP.
    Supports creating a new index or loading an existing one.
    Includes robustness improvements: DummyEmbeddings, dimension check, save path creation.
    (Logging removed for brevity).
    """
    top_k = settings.get("top_k", 5)
    user_requested_metric = settings.get("metric", "cosine").lower()
    faiss_mode = settings.get("faissMode", "create").lower()
    faiss_path = settings.get("faissPath", "") # Path to the FOLDER
    similarity_threshold = settings.get("similarity_threshold", 0) / 100

    # Basic Input Validation
    if not chunk_objs or not chunk_embeddings:
         # Consider raising an error or returning a structured error response
         return {q: [] for q in queries}
    if not queries or not query_embeddings:
         # Consider raising an error or returning a structured error response
         return {q: [] for q in queries}

    try:
        if not isinstance(chunk_embeddings[0], list) or not isinstance(query_embeddings[0], list):
             raise TypeError("Embeddings should be lists of lists of floats.")
        dimension = len(chunk_embeddings[0])
        query_dimension = len(query_embeddings[0])
    except (IndexError, TypeError) as e:
         # Consider raising an error or returning a structured error response
         # print(f"Error validating embedding structure: {e}") # Example direct print
         return {q: [] for q in queries}

    if dimension != query_dimension:
         # Consider raising an error or returning a structured error response
         # print(f"Embedding dimension mismatch: Chunks({dimension}), Queries({query_dimension})")
         return {q: [] for q in queries}

    chunk_embeddings_np = np.array(chunk_embeddings).astype('float32')
    query_embeddings_np = np.array(query_embeddings).astype('float32')

    try:
        dummy_embeddings = DummyEmbeddings(dimension=dimension)
    except ValueError as e:
         # print(f"Failed to initialize DummyEmbeddings: {e}")
         return {q: [] for q in queries}

    vector_store = None
    actual_metric = None

    # === Step 1: Initialize LangChain FAISS Vector Store ===
    if faiss_mode == "load":
        index_file = os.path.join(faiss_path, "index.faiss")
        pkl_file = os.path.join(faiss_path, "index.pkl")
        if not faiss_path or not os.path.exists(index_file) or not os.path.exists(pkl_file):
            # print(f"FAISS index not found in folder '{faiss_path}'")
            return {q: [] for q in queries}

        try:
            vector_store = FAISS.load_local(
                folder_path=faiss_path,
                embeddings=dummy_embeddings,
                index_name="index",
                allow_dangerous_deserialization=True
            )

            # Check Loaded Index Dimension
            loaded_dimension = vector_store.index.d
            if loaded_dimension != dimension:
                 # print(f"Dimension mismatch: Loaded index({loaded_dimension}), Provided queries({dimension})")
                 return {q: [] for q in queries}

            # Determine the metric of the LOADED index
            loaded_metric_type = vector_store.index.metric_type
            if loaded_metric_type == faiss.METRIC_L2:
                actual_metric = "l2"
            elif loaded_metric_type == faiss.METRIC_INNER_PRODUCT:
                actual_metric = "cosine"
            else:
                actual_metric = "l2" # Default fallback

            # Informational check (no logger)
            # if actual_metric != user_requested_metric:
            #      print(f"Note: Loaded metric ('{actual_metric}') differs from requested ('{user_requested_metric}'). Using '{actual_metric}'.")

        except Exception as e:
            # print(f"Error loading FAISS index from {faiss_path}: {e}")
            return {q: [] for q in queries}

    elif faiss_mode == "create":
        texts = [chunk.get("text", "") for chunk in chunk_objs]
        metadatas = [{"docTitle": chunk.get("docTitle", ""), "chunkId": chunk.get("chunkId", str(i))} for i, chunk in enumerate(chunk_objs)]

        try:
            if user_requested_metric in ["cosine", "ip"]:
                actual_metric = "cosine"
                faiss.normalize_L2(chunk_embeddings_np)
                index = faiss.IndexFlatIP(dimension)
            elif user_requested_metric == "l2":
                actual_metric = "l2"
                index = faiss.IndexFlatL2(dimension)
            else:
                actual_metric = "l2"
                index = faiss.IndexFlatL2(dimension)

            docstore = InMemoryDocstore({str(i): Document(page_content=texts[i], metadata=metadatas[i]) for i in range(len(texts))})
            index_to_docstore_id = {i: str(i) for i in range(len(texts))}

            index.add(chunk_embeddings_np)

            vector_store = FAISS(
                embedding_function=dummy_embeddings,
                index=index,
                docstore=docstore,
                index_to_docstore_id=index_to_docstore_id
            )

        except Exception as e:
             # print(f"Error during index creation process: {e}")
             return {q: [] for q in queries}

        if faiss_path:
            try:
                if not os.path.isdir(faiss_path):
                     os.makedirs(faiss_path, exist_ok=True)
                vector_store.save_local(folder_path=faiss_path, index_name="index")
            except Exception as e:
                 # print(f"Error saving FAISS index to {faiss_path}: {e}")
                 # print("Continuing retrieval despite save failure.")
                 pass # Continue even if saving fails

    else:
         # print(f"Invalid faissMode: '{faiss_mode}'. Must be 'create' or 'load'.")
         return {q: [] for q in queries}

    if not vector_store or not isinstance(vector_store, FAISS) or not actual_metric:
         # print("Failed to initialize a valid FAISS vector store object.")
         return {q: [] for q in queries}

    # === Step 2: Perform FAISS Retrieval ===
    results = {}
    for query, q_embedding in zip(queries, query_embeddings_np):
        try:
            query_vec = q_embedding.reshape(1, -1).astype('float32')
            if actual_metric == "cosine":
                faiss.normalize_L2(query_vec)

            search_results = vector_store.similarity_search_with_score_by_vector(
                embedding=query_vec[0],
                k=top_k
            )

            # === Step 3: Convert results, apply threshold, and format ===
            retrieved = []
            for doc, score in search_results:
                similarity_score = 0.0
                if actual_metric == "l2":
                    raw_score_l2 = float(score)
                    similarity_score = 1.0 / (1.0 + raw_score_l2) if raw_score_l2 >= 0 else 1.0
                elif actual_metric == "cosine":
                    raw_score_cosine = float(score)
                    similarity_score = max(0.0, min(1.0, raw_score_cosine))
                # else: fallback similarity_score remains 0.0

                if similarity_score >= similarity_threshold:
                    retrieved.append({
                        "text": doc.page_content,
                        "similarity": round(similarity_score, 6),
                        "docTitle": doc.metadata.get("docTitle", ""),
                        "chunkId": doc.metadata.get("chunkId", ""),
                    })

            retrieved.sort(key=lambda x: x["similarity"], reverse=True)
            results[query] = retrieved

        except Exception as e:
            # print(f"Error during similarity search for query '{query[:70]}...': {e}")
            results[query] = []

    return results

@RetrievalMethodRegistry.register("pinecone")
def handle_pinecone(chunk_objs, chunk_embeddings, queries, query_embeddings, settings):
    """
    Retrieve chunks using Pinecone with adaptable behavior based on mode: create, load.
    Includes smarter waiting based on index stats polling.
    """
    print("[DEBUG] Entered handle_pinecone function.")

    from pinecone import Pinecone, ServerlessSpec
    import uuid
    import time

    # 1. Extract settings
    top_k = settings.get("top_k", 5)
    # Fix: Handle score interpretation based on metric during retrieval
    similarity_function = settings.get("pineconeSimilarity", "cosine").lower() # Ensure lower case
    # Fix: Make threshold make sense for metric, or maybe remove for MVP?
    # Let's make threshold interpretation dependent on metric later in the code.
    # Raw threshold value from settings:
    raw_similarity_threshold = settings.get("similarity_threshold", 0) # Keep as is from settings for now

    pinecone_api_key = settings.get("pineconeApiKey", "")
    pinecone_env = settings.get("pineconeEnvironment", "us-east-1")
    pinecone_index_name = settings.get("pineconeIndex", "")
    pinecone_namespace = settings.get("pineconeNamespace", "")  # optional
    pinecone_mode = settings.get("pineconeMode", "create")  # "create", "load"

    # --- Smarter Wait Settings ---
    polling_interval_seconds = settings.get("pineconePollingInterval", 3) # Check every 3 seconds
    max_wait_time_seconds = settings.get("pineconeMaxWaitTime", 120) # Max wait 2 minutes

    print(f"[DEBUG] Pinecone settings extracted:")
    print(f"  top_k = {top_k}")
    print(f"  raw_similarity_threshold = {raw_similarity_threshold}") # Log raw value
    print(f"  pinecone_api_key = {'(HIDDEN)' if pinecone_api_key else '(MISSING)'}")
    print(f"  pinecone_env = {pinecone_env}")
    print(f"  pinecone_index_name = {pinecone_index_name}")
    print(f"  pinecone_namespace = {pinecone_namespace}")
    print(f"  similarity_function = {similarity_function}") # This IS used for index creation
    print(f"  pinecone_mode = {pinecone_mode}")
    print(f"  polling_interval = {polling_interval_seconds}s, max_wait_time = {max_wait_time_seconds}s")


    if not pinecone_api_key or not pinecone_index_name:
        print("[ERROR] Missing Pinecone API key or index name. Aborting.")
        return {q: [] for q in queries}
    if not chunk_embeddings: # Check for empty embeddings early
         print("[ERROR] chunk_embeddings list is empty. Aborting.")
         return {q: [] for q in queries}

    # 2. Initialize Pinecone
    print("[DEBUG] Initializing Pinecone client...")
    pc = Pinecone(api_key=pinecone_api_key)

    # 3. Check/Create Index
    print("[DEBUG] Checking existing Pinecone indexes...")
    existing_indexes_info = pc.list_indexes()
    existing_index_names = [idx["name"] for idx in existing_indexes_info]
    index_exists = pinecone_index_name in existing_index_names

    dimension = len(chunk_embeddings[0]) # Get dimension safely now

    index = None # Initialize index variable

    if pinecone_mode == "create":
        if index_exists:
            print(f"[DEBUG] Deleting existing Pinecone index '{pinecone_index_name}'...")
            try:
                pc.delete_index(pinecone_index_name)
                 # Wait a bit after delete before creating again
                print("[DEBUG] Waiting briefly after delete...")
                time.sleep(5)
            except Exception as e:
                 print(f"[WARN] Failed to delete index '{pinecone_index_name}': {e}. Trying to create anyway.")

        print(f"[DEBUG] Creating Pinecone index '{pinecone_index_name}'...")
        try:
            pc.create_index(
                name=pinecone_index_name,
                dimension=dimension,
                metric=similarity_function, # Correctly uses user's metric choice
                spec=ServerlessSpec(cloud="aws", region=pinecone_env)
            )
            # Wait a moment for index to initialize after creation
            print("[DEBUG] Index creation initiated. Waiting briefly...")
            time.sleep(10)
            print(f"[DEBUG] Index '{pinecone_index_name}' created/ready.")
            index = pc.Index(name=pinecone_index_name)
            upsert_chunks = True
        except Exception as e:
            print(f"[ERROR] Failed to create index '{pinecone_index_name}': {e}. Aborting.")
            return {q: [] for q in queries}

    elif pinecone_mode == "load":
        if not index_exists:
            print(f"[ERROR] Index '{pinecone_index_name}' does not exist. Cannot load. Aborting.")
            return {q: [] for q in queries}
        print(f"[DEBUG] Connecting to existing Pinecone index '{pinecone_index_name}'...")
        index = pc.Index(name=pinecone_index_name)
        # Check index dimension matches data
        try:
             stats = index.describe_index_stats()
             if stats.dimension != dimension:
                  print(f"[ERROR] Dimension mismatch: Index '{pinecone_index_name}' has dimension {stats.dimension}, but provided data has dimension {dimension}. Aborting.")
                  return {q: [] for q in queries}
             print(f"[DEBUG] Connected. Index dimension {stats.dimension} matches data.")
        except Exception as e:
             print(f"[WARN] Could not verify index dimension for '{pinecone_index_name}': {e}")
        upsert_chunks = True # Still upserting in load mode as per original logic
                           # Consider adding a 'connect_only' mode if needed

    else:
         print(f"[ERROR] Invalid pineconeMode '{pinecone_mode}'. Use 'create' or 'load'. Aborting.")
         return {q: [] for q in queries}

    # Upsert embeddings if needed
    if upsert_chunks and index:
        print("[DEBUG] Preparing vectors to upsert...")
        vectors_to_upsert = []
        for chunk, embedding in zip(chunk_objs, chunk_embeddings):
            vector_id = chunk.get("chunkId", str(uuid.uuid4())) # Ensure unique IDs
            metadata = {
                "text": chunk.get("text", ""),
                "docTitle": chunk.get("docTitle", ""),
                "chunkId": vector_id, # Store the ID in metadata too if needed
            }
            # Ensure embedding is a list of floats
            embedding_list = [float(x) for x in embedding]
            vectors_to_upsert.append((vector_id, embedding_list, metadata))

        if vectors_to_upsert:
            num_to_upsert = len(vectors_to_upsert)
            print(f"[DEBUG] Getting initial vector count in namespace '{pinecone_namespace}'...")
            initial_count = 0
            try:
                initial_stats = index.describe_index_stats()
                initial_count = initial_stats.namespaces.get(pinecone_namespace, {}).get('vector_count', 0)
                print(f"[DEBUG] Initial vector count: {initial_count}")
            except Exception as e:
                print(f"[WARN] Could not get initial vector count: {e}. Assuming 0.")

            expected_count = initial_count + num_to_upsert # Approximate target

            print(f"[DEBUG] Upserting {num_to_upsert} vectors into namespace '{pinecone_namespace}'...")
            try:
                # Upsert in batches for larger datasets if necessary (Pinecone handles batching internally to some extent)
                # For very large upserts (> few MB), consider batching client-side too.
                index.upsert(vectors=vectors_to_upsert, namespace=pinecone_namespace)
                print("[DEBUG] Upsert call completed.")

                print(f"[DEBUG] Polling index stats (every {polling_interval_seconds}s, max {max_wait_time_seconds}s) waiting for vector count to reach >= {expected_count}...")
                start_time = time.time()
                while True:
                    elapsed_time = time.time() - start_time
                    if elapsed_time > max_wait_time_seconds:
                        print(f"[WARN] Max wait time ({max_wait_time_seconds}s) exceeded while polling index stats. Proceeding anyway.")
                        break
                    try:
                        current_stats = index.describe_index_stats()
                        current_count = current_stats.namespaces.get(pinecone_namespace, {}).get('vector_count', 0)
                        print(f"[DEBUG] Polling: Current count = {current_count}, Target = {expected_count}, Time elapsed = {elapsed_time:.1f}s")

                        # Check if count is sufficient. >= handles potential overwrites where count might not increase exactly by num_to_upsert
                        if current_count >= expected_count:
                            print(f"[DEBUG] Target vector count reached ({current_count} >= {expected_count}). Index likely ready.")
                            break
                        # Handle edge case: If initial count was already >= expected (e.g., re-upserting same data)
                        if current_count >= initial_count and num_to_upsert > 0 and initial_count == expected_count:
                             print("[DEBUG] Count hasn't changed but was non-zero initially. Assuming upsert modified existing vectors.")
                             break


                    except Exception as e:
                        print(f"[WARN] Error polling index stats: {e}. Retrying...")

                    time.sleep(polling_interval_seconds)
                # Add a tiny final buffer sleep just in case? Optional.
                # time.sleep(1)
                print("[DEBUG] Finished waiting/polling.")
                # --- End Smarter Wait Logic ---

            except Exception as e:
                 print(f"[ERROR] Failed during upsert or polling: {e}")
                 # Decide if you want to proceed or abort if upsert/wait fails
                 return {q: [] for q in queries}
        else:
            print("[DEBUG] No vectors prepared to upsert.")

    # Query / Retrieval
    print("[DEBUG] Starting retrieval for queries...")
    results = {}
    if not index:
        print("[ERROR] Index object is not valid. Cannot perform retrieval.")
        return {q: [] for q in queries}

    for query, query_emb in zip(queries, query_embeddings):
        query_short = query[:70] + "..." if len(query) > 70 else query
        print(f"[DEBUG] Processing query: '{query_short}'")
        try:
            # Ensure query embedding is a list of floats
            query_emb_list = [float(x) for x in query_emb]

            pinecone_response = index.query(
                namespace=pinecone_namespace,
                vector=query_emb_list,
                top_k=top_k,
                include_metadata=True
            )

            # --- Fix Score Interpretation and Thresholding ---
            retrieved_chunks = []
            print(f"[DEBUG] Raw Pinecone response matches for query '{query_short}': {len(pinecone_response.get('matches', []))}")
            for match in pinecone_response.get("matches", []):
                score = match["score"]
                metadata = match.get("metadata", {})
                chunk_text = metadata.get("text", "")
                doc_title = metadata.get("docTitle", "")
                chunk_id = metadata.get("chunkId", match.id) # Fallback to match ID if not in metadata

                passes_threshold = False
                similarity_score_for_output = score # Default to raw score

                if similarity_function == "euclidean":
                     # Lower score (distance) is better. Threshold is max distance.
                     # Let's assume raw_similarity_threshold is meant as max distance here.
                     # A threshold of 0 for distance is impossible unless vectors are identical.
                     # Maybe skip thresholding for euclidean for MVP if threshold meaning is unclear?
                     # Or, convert distance to a similarity score first? 1 / (1 + distance)
                     similarity_score_for_output = 1.0 / (1.0 + score) if score >= 0 else 1.0
                     # Now apply the threshold (interpreted as min similarity) to the converted score
                     passes_threshold = (similarity_score_for_output >= (raw_similarity_threshold / 100.0))
                     print(f"  Match(euclidean): ID={chunk_id}, Dist={score:.4f}, Sim={similarity_score_for_output:.4f}, Threshold={raw_similarity_threshold/100.0}, Passes={passes_threshold}")

                elif similarity_function == "cosine":
                     # Higher score is better. Threshold is min similarity.
                     similarity_score_for_output = max(0.0, min(1.0, score)) # Clamp [0, 1]
                     passes_threshold = (similarity_score_for_output >= (raw_similarity_threshold / 100.0))
                     print(f"  Match(cosine): ID={chunk_id}, Score={score:.4f}, Sim={similarity_score_for_output:.4f}, Threshold={raw_similarity_threshold/100.0}, Passes={passes_threshold}")

                elif similarity_function == "dotproduct":
                     # Higher score is better. Threshold could be anything.
                     # Dot product isn't easily normalized to [0,1]. Thresholding is tricky.
                     # For MVP, maybe just skip thresholding for dotproduct or require user to know appropriate range.
                     # Let's apply the raw threshold directly, assuming user knows the scale.
                     similarity_score_for_output = score # Keep raw score
                     passes_threshold = (score >= raw_similarity_threshold) # Compare raw score to raw threshold
                     print(f"  Match(dotproduct): ID={chunk_id}, Score={score:.4f}, RawThreshold={raw_similarity_threshold}, Passes={passes_threshold}")
                else:
                     # Unknown metric
                     print(f"[WARN] Unknown similarity function '{similarity_function}' encountered during score interpretation.")
                     passes_threshold = True # Default to passing if metric is unknown? Or False? Let's pass.

                if passes_threshold:
                     retrieved_chunks.append({
                         "text": chunk_text,
                         "docTitle": doc_title,
                         "chunkId": chunk_id,
                         "similarity": similarity_score_for_output # Use the potentially converted score
                     })
                # else: Skip results not meeting the threshold/distance criteria (already logged above)

            # Sort final results if needed (Pinecone usually returns sorted, but doesn't hurt)
            # Make sure sorting key matches the similarity interpretation
            sort_reverse = (similarity_function != "euclidean") # Sort descending for cosine/dotproduct, ascending for raw euclidean dist if you kept it
            # Since we convert euclidean to similarity, always sort descending by 'similarity' field
            retrieved_chunks.sort(key=lambda x: x["similarity"], reverse=True)

            results[query] = retrieved_chunks
            print(f"[DEBUG] Retrieved {len(retrieved_chunks)} chunks for query: '{query_short}' after filtering/thresholding.")
        except Exception as e:
            print(f"[ERROR] Failed to process query '{query_short}': {e}")
            results[query] = []

    print("[DEBUG] Retrieval completed.")
    return results

# Add near other imports at the top
import chromadb
from chromadb.utils import embedding_functions
import uuid # For generating fallback IDs

# Define the DummyEmbeddings class if it's not already globally available
# (Assuming it's defined similarly to the handle_faiss implementation)
from langchain_core.embeddings import Embeddings
from typing import List
class DummyEmbeddings(Embeddings):
    """
    A dummy embedding class implementing the LangChain Embeddings interface.
    Used when pre-computed embeddings are provided.
    Returns zero vectors of the specified dimension.
    """
    def __init__(self, dimension: int):
        self.dimension = dimension
        if not isinstance(dimension, int) or dimension <= 0:
             raise ValueError(f"Dimension must be a positive integer, got {dimension}")

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """Return zero vectors for a list of documents."""
        # print(f"[DummyEmbeddings] embed_documents called for {len(texts)} texts. Returning zeros.") # Debug print
        return [[0.0] * self.dimension for _ in texts]

    def embed_query(self, text: str) -> List[float]:
        """Return a single zero vector for a query."""
        # print(f"[DummyEmbeddings] embed_query called for text: '{text[:50]}...'. Returning zeros.") # Debug print
        return [0.0] * self.dimension

@RetrievalMethodRegistry.register("chromadb")
def handle_chromadb(chunk_objs, chunk_embeddings, queries, query_embeddings, settings):
    """
    Retrieve chunks using Chroma DB Vector Store with precomputed embeddings.
    Supports in-memory or persistent storage and cosine/l2 metrics.
    """
    print("[ChromaDB] Starting retrieval with Chroma...", flush=True)

    # === Step 1: Extract Settings ===
    chroma_mode = settings.get("chromaMode", "memory").lower()  # "memory" or "persistent"
    chroma_path = settings.get("chromaPersistDir", "./chroma_db") # Default path if persistent
    collection_name = settings.get("chromaCollection", f"collection_{uuid.uuid4().hex[:8]}") # Default unique name
    # Chroma default is L2, let's allow specifying, default to cosine for consistency
    # Chroma uses 'l2', 'cosine', 'ip' (inner product)
    distance_metric = settings.get("chromaDistanceMetric", "cosine").lower()
    top_k = settings.get("top_k", 5)
    # similarity_threshold is 0-100 in settings, convert to 0-1
    similarity_threshold = settings.get("similarity_threshold", 0) / 100.0
    # Add a setting to control cleanup on create mode
    cleanup_on_create = settings.get("chromaCleanupOnCreate", True)

    print(f"[ChromaDB] Mode: {chroma_mode}", flush=True)
    print(f"[ChromaDB] Top K: {top_k}", flush=True)
    print(f"[ChromaDB] Similarity threshold: {similarity_threshold}", flush=True)
    print(f"[ChromaDB] Collection: {collection_name}", flush=True)
    print(f"[ChromaDB] Metric: {distance_metric}", flush=True)
    if chroma_mode == "persistent":
        print(f"[ChromaDB] Persistence path: {chroma_path}", flush=True)
    print(f"[ChromaDB] Cleanup on create: {cleanup_on_create}", flush=True)

    # === Basic Input Validation ===
    if not chunk_objs or not chunk_embeddings:
         print("[ChromaDB ERROR] No chunk objects or chunk embeddings provided.", flush=True)
         return {q: [] for q in queries}
    if not queries or not query_embeddings:
         print("[ChromaDB ERROR] No queries or query embeddings provided.", flush=True)
         return {q: [] for q in queries}

    try:
        if not isinstance(chunk_embeddings[0], list) or not isinstance(query_embeddings[0], list):
             raise TypeError("Embeddings should be lists of lists of floats.")
        dimension = len(chunk_embeddings[0])
        query_dimension = len(query_embeddings[0])
        if dimension == 0:
            raise ValueError("Chunk embedding dimension cannot be zero.")
        if query_dimension == 0:
             raise ValueError("Query embedding dimension cannot be zero.")
    except (IndexError, TypeError, ValueError) as e:
         print(f"[ChromaDB ERROR] Invalid embedding structure or dimension: {e}", flush=True)
         return {q: [] for q in queries}

    if dimension != query_dimension:
         print(f"[ChromaDB ERROR] Embedding dimension mismatch: Chunks({dimension}), Queries({query_dimension})", flush=True)
         return {q: [] for q in queries}

    # === Step 2: Initialize Chroma Client ===
    try:
        print(f"[ChromaDB] Initializing Chroma client (Mode: {chroma_mode})...", flush=True)
        if chroma_mode == "persistent":
            # Ensure the directory exists for persistent client
            if not os.path.exists(chroma_path):
                 print(f"[ChromaDB] Creating persistence directory: {chroma_path}", flush=True)
                 os.makedirs(chroma_path, exist_ok=True)
            chroma_client = chromadb.PersistentClient(path=chroma_path)
        else: # memory mode
            chroma_client = chromadb.Client()
        print("[ChromaDB] Chroma client initialized.", flush=True)

        # === Step 3: Get or Create Collection ===
        print(f"[ChromaDB] Accessing collection: '{collection_name}'", flush=True)

        # Handle cleanup if in create mode and collection exists
        if chroma_mode == "create" and cleanup_on_create:
             try:
                 existing_collections = [col.name for col in chroma_client.list_collections()]
                 if collection_name in existing_collections:
                      print(f"[ChromaDB] 'create' mode: Deleting existing collection '{collection_name}'...", flush=True)
                      chroma_client.delete_collection(name=collection_name)
                      print(f"[ChromaDB] Collection '{collection_name}' deleted.", flush=True)
                 else:
                      print(f"[ChromaDB] 'create' mode: Collection '{collection_name}' does not exist, no need to delete.", flush=True)
             except Exception as e:
                  print(f"[ChromaDB WARN] Failed to delete collection '{collection_name}' during cleanup: {e}. Proceeding...", flush=True)


        # Define metric for the collection (only possible at creation)
        # Use Chroma's constants for metadata keys if available, otherwise string
        # Note: distance_metric from settings should match Chroma's options ('l2', 'cosine', 'ip')
        if distance_metric not in ['l2', 'cosine', 'ip']:
            print(f"[ChromaDB WARN] Invalid distance metric '{distance_metric}'. Defaulting to 'cosine'.", flush=True)
            distance_metric = 'cosine'

        collection_metadata = {"hnsw:space": distance_metric}

        print(f"[ChromaDB] Getting or creating collection '{collection_name}' with metric '{distance_metric}'...", flush=True)
        collection = chroma_client.get_or_create_collection(
            name=collection_name,
            metadata=collection_metadata # Set metric here
            # embedding_function=embedding_functions.DefaultEmbeddingFunction() # Optional: Chroma can manage its own embeddings if needed, but we provide them
        )
        print(f"[ChromaDB] Collection '{collection_name}' ready.", flush=True)

        # === Step 4: Add/Upsert Data ===
        # Decide whether to add or upsert based on mode or specific setting if added later
        # For now, let's use upsert as it's generally safer and handles both create/load scenarios.
        print(f"[ChromaDB] Preparing {len(chunk_objs)} items for upsert...", flush=True)
        ids = []
        embeddings_to_add = []
        metadatas_to_add = []
        documents_to_add = [] # Chroma requires text content ('documents')

        for i, chunk in enumerate(chunk_objs):
            chunk_id = chunk.get("chunkId")
            # Ensure a unique, non-empty string ID
            if not chunk_id or not isinstance(chunk_id, str) or len(chunk_id.strip()) == 0:
                 chunk_id = f"chunk_{i}_{uuid.uuid4().hex[:8]}"
                 # print(f"[ChromaDB WARN] Generated fallback ID: {chunk_id} for chunk index {i}", flush=True)

            ids.append(chunk_id)
            # Ensure embeddings are lists of standard floats
            embeddings_to_add.append([float(e) for e in chunk_embeddings[i]])
            metadatas_to_add.append({
                "docTitle": chunk.get("docTitle", ""),
                "chunkId": chunk_id # Store original/generated ID in metadata too
                # Add any other relevant metadata from chunk_obj if needed
            })
            documents_to_add.append(chunk.get("text", "")) # Add the text content

        if ids:
            print(f"[ChromaDB] Upserting {len(ids)} items into collection '{collection_name}'...", flush=True)
            # Upsert handles adding new and updating existing IDs
            collection.upsert(
                ids=ids,
                embeddings=embeddings_to_add,
                metadatas=metadatas_to_add,
                documents=documents_to_add
            )
            print("[ChromaDB] Upsert operation completed.", flush=True)
            # Optional: Add a small delay if needed, though upsert is usually synchronous
            # time.sleep(1)
        else:
             print("[ChromaDB] No items to upsert.", flush=True)


        # === Step 5: Perform Retrieval ===
        print("[ChromaDB] Starting retrieval for queries...", flush=True)
        results = {}

        for query, q_embedding in zip(queries, query_embeddings):
            query_short = query[:70] + "..." if len(query) > 70 else query
            print(f"[ChromaDB] Processing query: '{query_short}'", flush=True)

            try:
                # Ensure query embedding is list of floats
                query_embedding_float = [float(e) for e in q_embedding]

                # Use collection.query
                query_results = collection.query(
                    query_embeddings=[query_embedding_float], # Query expects a list of embeddings
                    n_results=top_k,
                    include=['metadatas', 'documents', 'distances'] # Request distances
                )

                # === Step 6: Process and Format Results ===
                retrieved = []
                # query_results structure: {'ids': [[id1,..]], 'documents': [[doc1,..]], 'metadatas': [[meta1,..]], 'distances': [[dist1,..]]}
                # Since we query one vector at a time, we access the first element of each list
                if query_results and query_results.get('ids') and query_results['ids'][0]:
                    num_results = len(query_results['ids'][0])
                    print(f"[ChromaDB] Query returned {num_results} raw results.", flush=True)

                    for i in range(num_results):
                        distance = query_results['distances'][0][i]
                        metadata = query_results['metadatas'][0][i]
                        doc_text = query_results['documents'][0][i]
                        result_id = query_results['ids'][0][i]

                        # Convert distance to similarity score (0 to 1, higher is better)
                        similarity_score = 0.0
                        if distance_metric == 'cosine':
                            # Cosine distance = 1 - cosine similarity
                            # similarity = 1 - distance
                            # Clamp between 0 and 1, as distance can theoretically be > 1 slightly due to float errors, or < 0 if metric='ip'
                            similarity_score = max(0.0, min(1.0, 1.0 - float(distance)))
                        elif distance_metric == 'l2':
                            # L2 distance is >= 0. Lower is better.
                            # Convert to similarity: 1 / (1 + distance) -> maps [0, inf) to (0, 1]
                            similarity_score = 1.0 / (1.0 + float(distance))
                        elif distance_metric == 'ip':
                             # Inner product: higher is better. Can be negative.
                             # No standard way to normalize IP to [0, 1] without knowing the range.
                             # Let's pass the raw score, but warn user. Or maybe try simple scaling?
                             # For now, just clamp positive values? Or return raw? Let's clamp [0,1] after possible normalization.
                             # Simple approach: If embeddings are normalized, IP is cosine sim. Assume user normalized if using IP.
                             similarity_score = max(0.0, min(1.0, float(distance))) # Treat like cosine if normalized
                             # print(f"[ChromaDB WARN] Inner Product metric used. Assuming normalized embeddings. Score={distance}, Similarity={similarity_score}", flush=True)

                        # Apply threshold
                        if similarity_score >= similarity_threshold:
                            retrieved.append({
                                "text": doc_text,
                                "similarity": round(similarity_score, 6), # Use calculated similarity
                                "docTitle": metadata.get("docTitle", ""),
                                # Use chunkId from metadata if available, else the result ID
                                "chunkId": metadata.get("chunkId", result_id),
                                # Optionally include raw distance for debugging:
                                # "raw_distance": float(distance),
                                # "metric": distance_metric
                            })
                        else:
                            print(f"[ChromaDB] Result {result_id} skipped due to threshold (Sim: {similarity_score:.4f} < Threshold: {similarity_threshold:.4f})", flush=True)

                    # Sort by similarity score descending (highest first)
                    retrieved.sort(key=lambda x: x["similarity"], reverse=True)

                    # Ensure we only return top_k results *after* filtering
                    retrieved = retrieved[:top_k]

                    results[query] = retrieved
                    print(f"[ChromaDB] Retrieved {len(retrieved)} chunks for query '{query_short}' after filtering/sorting.", flush=True)
                else:
                    print(f"[ChromaDB] No results returned from Chroma query for '{query_short}'.", flush=True)
                    results[query] = []

            except Exception as e:
                print(f"[ChromaDB ERROR] Failed during query or result processing for query '{query_short}': {e}", flush=True)
                # Add traceback for debugging if needed
                # import traceback
                # traceback.print_exc()
                results[query] = [] # Return empty list for this query on error

    except Exception as e:
        print(f"[ChromaDB FATAL ERROR] Failed during client/collection setup or upsert: {e}", flush=True)
        # import traceback
        # traceback.print_exc()
        # Return empty results for all queries if setup failed
        results = {q: [] for q in queries}

    print("[ChromaDB] Retrieval process finished.", flush=True)
    return results

from azure.cosmos import CosmosClient, PartitionKey
import numpy as np
import uuid
import time

@RetrievalMethodRegistry.register("cosmosdb")
def handle_cosmosdb(chunk_objs, chunk_embeddings, queries, query_embeddings, settings):
    """
    Retrieve chunks using Cosmos DB Vector Store.

    Args:
        chunk_objs (list): List of dicts with chunk data.
        chunk_embeddings (list): List of chunk embedding vectors.
        queries (list): List of query strings.
        query_embeddings (list): List of query embedding vectors.
        settings (dict): Configuration settings.

    Returns:
        dict: Mapping of query strings to retrieved results.
    """

    # === Extract settings ===
    cosmos_mode = settings.get("cosmosMode", "create")
    endpoint = settings.get("cosmosEndpoint")
    key = settings.get("cosmosKey")
    database_name = settings.get("cosmosDatabase", "VectorDB")
    container_name = settings.get("cosmosContainer", "Chunks")
    top_k = settings.get("top_k", 5)
    similarity_threshold = settings.get("similarity_threshold", 0) / 100

    # === Step 1: Init Cosmos Client ===
    print("[CosmosDB] Connecting to Cosmos DB...")
    client = CosmosClient(endpoint, credential=key)
    db = client.create_database_if_not_exists(id=database_name)
    container = db.create_container_if_not_exists(
        id=container_name,
        partition_key=PartitionKey(path="/chunkId"),
        offer_throughput=400
    )

    dimension = len(chunk_embeddings[0])
    print(f"[CosmosDB] Vector dimension: {dimension}")
    print(f"[CosmosDB] Mode: {cosmos_mode}")

    # === Step 2: Create or Load Embeddings ===
    if cosmos_mode == "create":
        print(f"[CosmosDB] Inserting {len(chunk_objs)} vectors...")
        for chunk, embedding in zip(chunk_objs, chunk_embeddings):
            try:
                item = {
                    "id": chunk.get("chunkId", str(uuid.uuid4())),
                    "chunkId": chunk.get("chunkId", str(uuid.uuid4())),
                    "text": chunk.get("text", ""),
                    "docTitle": chunk.get("docTitle", ""),
                    "embedding": embedding
                }
                container.upsert_item(item)
            except Exception as e:
                print(f"[CosmosDB ERROR] Failed to upsert item: {e}")
        print("[CosmosDB] Upsert complete. Waiting 10s for indexing...")
        time.sleep(10)

    # === Step 3: Perform Retrieval ===
    print("[CosmosDB] Starting retrieval...")
    results = {}

    for query, q_embedding in zip(queries, query_embeddings):
        print(f"[CosmosDB] Query: '{query}'")

        vector_query = {
            "vector": q_embedding,
            "topK": top_k,
            "fields": "embedding",
            "numCandidates": 100,
            "metric": "cosine"
        }

        try:
            vector_sql = {
                "query": "SELECT * FROM c WHERE IS_VECTOR_SEARCHABLE(c.embedding)",
                "vectorSearch": vector_query
            }

            response = container.query_items(
                query=vector_sql,
                enable_cross_partition_query=True
            )

            retrieved = []
            for item in response:
                score = item.get("_vectorScore", 0)
                if score >= similarity_threshold:
                    retrieved.append({
                        "text": item.get("text", ""),
                        "docTitle": item.get("docTitle", ""),
                        "chunkId": item.get("chunkId", ""),
                        "similarity": float(score)
                    })

            results[query] = retrieved
            print(f"[CosmosDB] Retrieved {len(retrieved)} chunks for query.")

        except Exception as e:
            print(f"[CosmosDB ERROR] Failed on query '{query}': {e}")
            results[query] = []

    print("[CosmosDB] Retrieval complete.")
    return results


@app.route("/retrieve", methods=["POST"])
def retrieve():
    """
    Process multiple retrieval methods against provided chunks and queries.
    
    Expected request format:
    {
        "methods": [
            {
                "id": "unique_method_id",
                "baseMethod": "bm25",
                "methodName": "BM25",
                "library": "BM25",
                "settings": { "top_k": 5, ... }
            },
            {
                "id": "unique_method_id",
                "baseMethod": "cosine",
                "methodName": "Cosine similarity",
                "library": "Cosine",
                "embeddingProvider": "HuggingFace Transformers",
                "settings": { "embeddingModel": "all-MiniLM-L6-v2", "top_k": 5, ... }
            }
            ...
        ],
        "chunks": [{"text": "...", "docTitle": "...", "chunkId": "..."}, ...],
        "queries": ["query1", "query2", ...] 
    }

    Expected output format:
    {
        "results": {
            "method_id": {
                "baseMethod": "bm25",
                "methodName": "BM25",
                "retrieved": {"query1": [chunk1, chunk2, ...], "query2": [...], ...},
                "status": "success"
            },
            ...
        }
    }
    
    Returns progressive results as each method completes.
    """
    data = request.json
    methods = data.get("methods", [])
    chunks = data.get("chunks", [])
    queries = data.get("queries", [])
    print(methods)
    # Validate inputs
    if not methods:
        return jsonify({"error": "No retrieval methods provided"}), 400
    if not chunks:
        return jsonify({"error": "No chunks provided"}), 400
    if not queries:
        return jsonify({"error": "No queries provided"}), 400
    

    # Group methods by embedding model to avoid redundant embedding computation
    embedding_methods = {}  # model -> list of methods requiring this model
    keyword_methods = []    # methods not requiring embeddings

    for method in methods:
        try:
            embedding_provider= method.get("embeddingProvider", None)
            if embedding_provider:
                # This is an embedding-based method
                embedding_model = method.get("settings")["embeddingModel"]
                full_embedder = f"{embedding_provider}#{embedding_model}"    
                if full_embedder not in embedding_methods:
                    embedding_methods[full_embedder] = []
                embedding_methods[full_embedder].append(method)
            else:
                # Non-embedding method
                keyword_methods.append(method)
        except Exception as e:
            results[method_id] = {
                "baseMethod": base_method,
                "methodName": method.get("methodName"),
                "error": str(e),
                "status": "error"
            }
            
    
    # Prepare the response object
    results = {}

    # Process keyword methods
    for method in keyword_methods:
        method_id = method.get("id")
        base_method = method.get("baseMethod")
        
        try:
            handler = RetrievalMethodRegistry.get_handler(base_method)
            if not handler:
                raise ValueError(f"Unknown method: {base_method}")
            retrieved = handler(chunks, queries, method.get("settings", {}))
                
            results[method_id] = {
                "baseMethod": base_method,
                "methodName": method.get("methodName"),
                "retrieved": retrieved,
                "status": "success"
            }
        except Exception as e:
            # Handle errors
            results[method_id] = {
                "baseMethod": base_method,
                "methodName": method.get("methodName"),
                "error": str(e),
                "status": "error"
            }

    # Process embedding-based methods grouped by model
    for embedder, methods in embedding_methods.items():
        try:
            provider, model_name = embedder.split("#", 1)
            embedder = EmbeddingModelRegistry.get_embedder(provider)
            if not embedder:
                        raise ValueError(f"Unknown embedding model: {model_name}")
            
            # Compute embeddings once for all methods using this model
            chunk_embeddings = embedder([c["text"] for c in chunks])
            query_embeddings = embedder(queries)
        except Exception as e:
                # Mark all methods using this model as failed
                for method in methods:
                    results[method.get("id")] = {
                        "baseMethod": method.get("baseMethod"),
                        "methodName": method.get("methodName"),
                        "error": f"Embedding model error: {str(e)}",
                        "status": "error"
                    }
                continue
        
        # Process each method with the same embeddings
        for method in methods:
            method_id = method.get("id")
            base_method = method.get("baseMethod")
            
            try:
                handler = RetrievalMethodRegistry.get_handler(base_method)
                print(handler)
                if not handler:
                    raise ValueError(f"Unknown method: {base_method}")                    
                retrieved = handler(chunks, chunk_embeddings, queries, query_embeddings, method.get("settings", {}))
                    
                results[method_id] = {
                    "baseMethod": base_method,
                    "methodName": method.get("methodName"),
                    "retrieved": retrieved,
                    "status": "success",
                    "embeddingModel": model_name
                }
            except Exception as e:
                # Handle errors
                results[method_id] = {
                    "baseMethod": base_method,
                    "methodName": method.get("methodName"),
                    "error": str(e),
                    "status": "error",
                    "embeddingModel": model_name
                }
    
    return jsonify({"results": results}), 200

if __name__ == "__main__":
    app.run(debug=True, port="5001")